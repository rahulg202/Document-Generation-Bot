import streamlit as st
import fitz  # PyMuPDF for PDF handling
import docx  # python-docx for DOCX handling
import markdown
import re
from io import BytesIO

def extract_text_from_pdf(pdf_file):
    """Extract text from a PDF file."""
    text = ""
    try:
        with fitz.open(stream=pdf_file.read(), filetype="pdf") as doc:
            for page in doc:
                text += page.get_text()
        return text
    except Exception as e:
        st.error(f"Error extracting text from PDF: {str(e)}")
        return ""

def extract_text_from_docx(docx_file):
    """Extract text from a DOCX file."""
    text = ""
    try:
        doc = docx.Document(docx_file)
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text
    except Exception as e:
        st.error(f"Error extracting text from DOCX: {str(e)}")
        return ""

def extract_text_from_txt(txt_file):
    """Extract text from a TXT file."""
    try:
        return txt_file.getvalue().decode("utf-8")
    except Exception as e:
        st.error(f"Error extracting text from TXT: {str(e)}")
        return ""

def extract_text_from_uploaded_file(uploaded_file):
    """Extract text from an uploaded file based on its type."""
    if uploaded_file.type == "application/pdf":
        return extract_text_from_pdf(uploaded_file)
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return extract_text_from_docx(uploaded_file)
    elif uploaded_file.type == "text/plain":
        return extract_text_from_txt(uploaded_file)
    else:
        st.error(f"Unsupported file type: {uploaded_file.type}")
        return ""

def markdown_to_html(markdown_text):
    """Convert markdown to HTML with proper styling."""
    # First, remove any HTML+Jinja or CSS markers that might be at the start
    markdown_text = re.sub(r'^html\+jinja\s*', '', markdown_text)
    markdown_text = re.sub(r'^body\s*{.*?}\s*h1,\s*h2,\s*h3\s*{.*?}\s*table\s*{.*?}.*?$', '', markdown_text, flags=re.MULTILINE)
    
    html = markdown.markdown(markdown_text)
    
    # Add styling to make it look better
    styled_html = f"""
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Generated Document</title>
        <style>
            body {{
                font-family: Arial, sans-serif;
                line-height: 1.6;
                margin: 2cm;
                color: #333;
            }}
            h1, h2, h3, h4, h5, h6 {{
                color: #2c3e50;
                margin-top: 1.5em;
                margin-bottom: 0.5em;
            }}
            h1 {{ font-size: 2em; border-bottom: 1px solid #eee; padding-bottom: 0.3em; }}
            h2 {{ font-size: 1.5em; }}
            p {{ margin-bottom: 1em; }}
            ul, ol {{ padding-left: 2em; margin-bottom: 1em; }}
            li {{ margin-bottom: 0.5em; }}
            table {{ border-collapse: collapse; width: 100%; margin-bottom: 1em; }}
            th, td {{ padding: 8px; text-align: left; border: 1px solid #ddd; }}
            th {{ background-color: #f2f2f2; }}
            blockquote {{ background-color: #f9f9f9; border-left: 4px solid #ccc; margin: 1em 0; padding: 0.5em 1em; }}
        </style>
    </head>
    <body>
        {html}
    </body>
    </html>
    """
    return styled_html

def html_to_pdf(html_content):
    """Convert HTML to PDF using ReportLab."""
    try:
        # Import ReportLab libraries
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        
        # Simple HTML parser for ReportLab (very basic)
        # For complex HTML, a more robust solution would be needed
        def convert_html_to_reportlab(html_content):
            # Extract body content
            body_match = re.search(r'<body>(.*?)</body>', html_content, re.DOTALL)
            if body_match:
                body_content = body_match.group(1)
            else:
                body_content = html_content
            
            # Very simple HTML tag removal - this is a basic approach
            # For production, you'd want a proper HTML to ReportLab converter
            text = re.sub(r'<[^>]*>', '', body_content)
            text = re.sub(r'\n\s*\n', '\n\n', text)  # Clean up whitespace
            
            return text
        
        # Create a PDF buffer
        buffer = BytesIO()
        
        # Create the PDF document
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        
        # Create custom style
        normal_style = ParagraphStyle(
            'Normal',
            fontName='Helvetica',
            fontSize=10,
            leading=12,
            spaceAfter=6
        )
        
        # Build the PDF content
        content = []
        
        # Convert HTML to text for ReportLab
        text = convert_html_to_reportlab(html_content)
        
        # Split by paragraphs
        paragraphs = text.split('\n\n')
        
        for para in paragraphs:
            if para.strip():
                p = Paragraph(para, normal_style)
                content.append(p)
                content.append(Spacer(1, 0.2 * inch))
        
        # Build the PDF
        doc.build(content)
        
        # Get the PDF content
        pdf_content = buffer.getvalue()
        buffer.close()
        
        return pdf_content
        
    except Exception as e:
        st.error(f"Error converting to PDF: {str(e)}")
        st.info("Trying alternative PDF generation method...")
        
        try:
            # Fallback to a simpler method using FPDF
            from fpdf import FPDF
            
            # Basic HTML to text
            text = re.sub(r'<[^>]*>', '', html_content)
            
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            
            # Split text to fit on page
            for line in text.split('\n'):
                if line.strip():
                    pdf.multi_cell(0, 10, line)
                    pdf.ln(2)
            
            return pdf.output(dest='S').encode('latin1')
            
        except Exception as fallback_error:
            st.error(f"Fallback PDF generation failed: {str(fallback_error)}")
            st.warning("Please try downloading as DOCX instead.")
            return None

def html_to_docx(html_content):
    """Convert HTML to DOCX with better formatting."""
    try:
        # Create a new DOCX document
        doc = docx.Document()
        
        # Extract content from HTML (basic approach)
        # In a production environment, you'd want a more robust HTML to DOCX converter
        
        # Extract body content
        body_match = re.search(r'<body>(.*?)</body>', html_content, re.DOTALL)
        if body_match:
            body_content = body_match.group(1)
        else:
            body_content = html_content
        
        # Process headings (h1, h2, h3)
        # This is a simplified approach - a real solution would parse the DOM properly
        heading_pattern = re.compile(r'<h([1-3])>(.*?)</h\1>', re.DOTALL)
        current_position = 0
        
        for match in heading_pattern.finditer(body_content):
            # Add text before this heading
            text_before = body_content[current_position:match.start()]
            if text_before.strip():
                # Remove HTML tags and add as paragraph
                clean_text = re.sub(r'<[^>]*>', '', text_before)
                if clean_text.strip():
                    doc.add_paragraph(clean_text.strip())
            
            # Add the heading with appropriate style
            heading_level = int(match.group(1))
            heading_text = re.sub(r'<[^>]*>', '', match.group(2))
            
            if heading_level == 1:
                doc.add_heading(heading_text, level=1)
            elif heading_level == 2:
                doc.add_heading(heading_text, level=2)
            else:
                doc.add_heading(heading_text, level=3)
            
            current_position = match.end()
        
        # Add remaining text
        remaining_text = body_content[current_position:]
        if remaining_text.strip():
            # Remove HTML tags
            clean_text = re.sub(r'<[^>]*>', '', remaining_text)
            # Split into paragraphs based on line breaks
            paragraphs = re.split(r'<br\s*/?>|\n+', clean_text)
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para.strip())
        
        # Save the document to a BytesIO object
        docx_io = BytesIO()
        doc.save(docx_io)
        docx_io.seek(0)
        
        return docx_io.getvalue()
    except Exception as e:
        st.error(f"Error converting to DOCX: {str(e)}")
        return None